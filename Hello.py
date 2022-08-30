from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# name, phone mumber and email address
name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

#profile picture
document.add_picture('pro.jpeg', width=Inches(2.0))

document.add_paragraph(name + '|' + phone_number + '|' + email)

# about me
document.add_heading('About Me')

document.add_paragraph(input('Tell about yourself? '))

#work experience
document.add_heading('Work Experience:')

p = document.add_paragraph()

company = input('What is your company name? ')
from_date = input('From date ')
to_date = input('To date ')

p.add_run(company + ':').bold = True
p.add_run(from_date + '-' + to_date + ':' + '\n'). italic = True

experience_details = input('describe your experiences: ')
p.add_run(experience_details)

# more experiences

while True:
    has_more_experiences = input('Do you have mroe experiences? Yes or No  ')

    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('What is your company name? ')
        from_date = input('From date ')
        to_date = input('To date ')

        p.add_run(company + ':').bold = True
        p.add_run(from_date + '-' + to_date + ':' + '\n'). italic = True

        experience_details = input('describe your experiences: ')
        p.add_run(experience_details)

    else:
        break

# skills
document.add_heading('Skills:')

skill = input('Enter skills ')
p = document.add_paragraph()

p.style = 'List Bullet'

p.add_run(skill)

while True:
    has_more_skills = input('Do you have more skills?  Yes or No ')

    if has_more_skills.lower() == 'yes':
        skill = input('Enter skills ')
        p = document.add_paragraph()

        p.style = 'List Bullet'

        p.add_run(skill)

    else:
        break

document.save('cv.docx')